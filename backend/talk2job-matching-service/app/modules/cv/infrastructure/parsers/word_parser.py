from docx import Document
import os


def extract(file_path: str) -> str:
    try:
        if not os.path.exists(file_path):
            return "Erreur : Fichier introuvable."

        doc = Document(file_path)
        full_text = []

        # 1. Extraction par paragraphes
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text.strip())

        # 2. Extraction par tableaux (indispensable pour les CV Tech)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    # On récupère le texte de chaque paragraphe dans la cellule
                    texts = [p.text.strip() for p in cell.paragraphs if p.text.strip()]
                    full_text.extend(texts)

        # 3. Fallback : Si toujours vide, on fouille dans le XML brut
        if not full_text:
            for element in doc.element.body.iter():
                # w:t est la balise texte standard de Word
                if element.tag.endswith("t") and element.text:
                    full_text.append(element.text.strip())

        result = "\n".join(full_text)
        return (
            result if result.strip() else "Document Word vide ou format non supporté."
        )

    except Exception as e:
        return f"Erreur critique Word : {str(e)}"
